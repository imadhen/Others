from docx import Document
from docx.shared import RGBColor
import tkinter as tk
from tkinter import filedialog, Text, messagebox

#doc = Document("asset/NEM.docx")
doc = Document("asset/Pathologie de l_antéhypophyse.docx")

def process_qcms(doc):
    new_qcm = []
    qcm =[]
    Permission = False

    for para in doc.paragraphs:
        text = para.text
        
        for run in para.runs:  # Each 'run' is a section with the same style in the paragraph
            if run.bold: #Detecter les questions by the bold text
                if Permission == True:
                    new_qcm.append("\n /// \n")
                if Permission == False:
                    Permission = True
            
            if run.font.color and run.font.color.rgb == RGBColor(255, 0, 0):
                text += "  **"
        
        new_qcm.append(text) #Ajouter la pargraphe au nouveau qcm

    for p in new_qcm:
        print(p)
        
    return new_qcm

def save_file(output_text):
    # Demander à l'utilisateur où sauvegarder le fichier
    file_path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text files", "*.txt"), ("DOCX files", "*.docx")])
    if not file_path:
        return

    if file_path.endswith(".txt"):
        with open(file_path, "w", encoding="utf-8") as output_file:
            output_file.write(output_text)
        messagebox.showinfo("Succès", f"Le fichier a été sauvegardé sous '{file_path}'.")
    elif file_path.endswith(".docx"):
        # Sauvegarder en format DOCX
        doc = Document()
        for qcm in output_text:
            doc.add_paragraph(qcm)
        doc.save(file_path)
        messagebox.showinfo("Succès", f"Le fichier a été sauvegardé sous '{file_path}'.")

def open_file():
    filepath = filedialog.askopenfilename(filetypes=[("DOCX files", "*.docx")])
    if not filepath:
        return

    try:
        print(filepath)
        doc = Document(filepath)
        processed_qcms = process_qcms(doc)
        output_text = processed_qcms

        # Sauvegarde du fichier
        save_file(output_text)

    except Exception as e:
        messagebox.showerror("Erreur", f"Une erreur s'est produite : {e} {filepath}")

# Interface Tkinter

root = tk.Tk()
root.title("Détection de QCM")

canvas = tk.Canvas(root, height=400, width=400)
canvas.pack()

frame = tk.Frame(root, bg="white")
frame.place(relwidth=0.8, relheight=0.8, relx=0.1, rely=0.1)

open_file_button = tk.Button(frame, text="Ouvrir fichier DOCX", padx=10, pady=5, fg="white", bg="#263D42", command=open_file)
open_file_button.pack()

root.mainloop()
            
